from webscrap.webscrap_single import get_chat_model
from langchain.chains.llm import LLMChain
from langchain_core.prompts import PromptTemplate
import PyPDF2
from fpdf import FPDF
from docx import Document

chat_model=get_chat_model()

def enToCh():
    prompt_str = """
    你是一位学术语言翻译专家，知晓很多学术英语，擅长将英文翻译成中文。请将以下文本从英文翻译成中文,注意，只要翻译后的内容，不需要说其他的\n\n
    {text}"""
    prompt = PromptTemplate(
        input_variables=["text"],
        template=prompt_str,
    )
    chain = LLMChain(
        prompt=prompt,
        llm=chat_model,
    )
    return chain

trans = enToCh()

def trans_pdf(file_path,output_path):
    count = 0
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        pdf_text = ''
        for page in pdf_reader.pages:
            count += 1
            print(count / len(pdf_reader.pages))
            temp = trans.invoke({"text": page.extract_text()})["text"]
            pdf_text += temp
    # 实例PDF对象
    pdf = FPDF('P', 'mm', 'A4')  # 使用A4纸张尺寸
    # 添加一页
    pdf.add_page()

    # 注册并使用支持中文的字体，例如 simhei.ttf
    pdf.add_font('SimHei', '', './models/simkai.ttf')  # 确保路径正确
    pdf.set_font("SimHei", size=12)

    # 设置缩进的空格数
    indent_spaces = '  '  # 两个空格
    lines = pdf_text.split('\n')

    # 页面宽度
    page_width = pdf.w - 2 * pdf.l_margin

    # 设置行高
    line_height = 10  # 使用适中的行高

    # 写入每一段到 PDF
    for line in lines:
        if line.strip():  # 仅处理非空行
            # 在段首添加缩进
            line = indent_spaces + line
            # 使用 multi_cell 自动换行
            pdf.multi_cell(page_width, line_height, text=line, align='L')
            # 段落间增加空行
            pdf.ln(line_height / 2)  # 使用一半行高增加空行

    # 保存 PDF 文件
    response = pdf.output(output_path)
    return response

def trans_docx(file_path,output_path):
    document = Document(file_path)
    paragraphs = document.paragraphs
    pages = []
    temp = []
    for i, paragraph in enumerate(paragraphs):
        temp.append(paragraph.text)
        if (i + 1) % 4 == 0:
            pages.append("\n".join(temp))
            temp = []

    # Add the last page if it has remaining content
    if temp:
        pages.append("\n".join(temp))

    text=""
    for i,page in enumerate(pages):
        print(i/len(pages))
        temp = trans.invoke({"text": page})["text"]
        text+=temp
        print(temp)

    # 创建一个新的 Word 文档
    doc = Document()

    # 添加段落并写入字符串
    doc.add_paragraph(text)

    # 保存文档
    response=doc.save(output_path)
    return response


def trans_txt(file_path,output_path):
    # 读取文件内容
    with open(file_path, "r", encoding="utf-8") as file:
        content = file.read()

    paragraphs = content.split('\n')
    pages = []
    temp = []
    for i, paragraph in enumerate(paragraphs):
        temp.append(paragraph)
        if (i + 1) % 4 == 0:
            pages.append("\n".join(temp))
            temp = []

    # Add the last page if it has remaining content
    if temp:
        pages.append("\n".join(temp))
    text = ""
    for i, page in enumerate(pages):
        temp = trans.invoke({"text": page})["text"]
        text += temp

    with open(output_path, "w", encoding="utf-8") as file:
        file.write(text)


